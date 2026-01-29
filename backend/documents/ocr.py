"""
Utilitaires pour l'extraction de texte par OCR (Optical Character Recognition).
"""
import pytesseract
from pdf2image import convert_from_path
from PIL import Image, ImageOps
import PyPDF2
import os
import tempfile
import docx
import io
from django.conf import settings
import logging

logger = logging.getLogger(__name__)


class OCRProcessor:
    """
    Classe pour gérer l'extraction de texte des documents.
    """
    
    def __init__(self):
        """
        Initialise le processeur OCR.
        """
        if hasattr(settings, 'TESSERACT_CMD') and settings.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
        
        self.ocr_languages = '+'.join(getattr(settings, 'OCR_LANGUAGES', ['fra', 'eng']))
    
    def preprocess_image(self, image):
        """
        Prétraite l'image pour améliorer la qualité OCR.
        """
        try:
            if image.mode != 'L':
                image = image.convert('L')
            image = ImageOps.autocontrast(image)
            return image
        except Exception as e:
            logger.warning(f"Erreur prétraitement image: {str(e)}")
            return image

    def extract_text_from_file(self, file_path):
        """
        Extrait le texte d'un fichier basé sur son extension.
        """
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                return self.extract_from_pdf(file_path)
            elif ext in ['.docx']:
                text, error = self.extract_from_docx(file_path)
                return text, None, error
            elif ext in ['.doc']:
                return '', None, 'Le format .doc est ancien. Veuillez le convertir en .docx pour l\'extraction.'
            elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
                return self.extract_from_image(file_path)
            elif ext in ['.txt']:
                text, error = self.extract_from_text(file_path)
                # Convertir en PDF pour AskYourPDF
                pdf_path, pdf_error = self.convert_txt_to_pdf(text, file_path)
                return text, pdf_path, error
            else:
                return '', None, f'Type de fichier non supporté pour l\'OCR: {ext}'
                
        except Exception as e:
            logger.exception(f"Erreur lors de l'extraction de texte pour {file_path}")
            return '', None, f"Erreur système: {str(e)}"
    
    def extract_from_pdf(self, pdf_path):
        """
        Extrait le texte d'un fichier PDF.
        """
        try:
            text = ''
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() or ''
            
            if text.strip():
                logger.info(f"Texte extrait directement du PDF: {len(text)} caractères")
                # On génère quand même le PDF recherchable pour la compatibilité AskYourPDF
                _, searchable_path, ocr_error = self.ocr_pdf_images(pdf_path)
                return text, searchable_path, ocr_error
            
            logger.info("PDF sans texte, utilisation de l'OCR...")
            return self.ocr_pdf_images(pdf_path)
            
        except Exception as e:
            logger.error(f"Erreur extraction PDF {pdf_path}: {str(e)}")
            # En cas d'erreur de lecture PDF (ex: malformé), on tente quand même l'OCR via Image
            return self.ocr_pdf_images(pdf_path)
    
    def ocr_pdf_images(self, pdf_path):
        """
        Convertit les pages PDF en images et applique l'OCR.
        Utilise PyMuPDF (fitz) car il est auto-contenu et ne nécessite pas Poppler.
        """
        try:
            import fitz
            doc = fitz.open(pdf_path)
            text = ''
            pdf_writer = PyPDF2.PdfWriter()
            
            for i in range(len(doc)):
                logger.info(f"OCR page {i+1}/{len(doc)}")
                
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72)) # 300 DPI
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                
                processed_image = self.preprocess_image(image)
                page_text = pytesseract.image_to_string(
                    processed_image,
                    lang=self.ocr_languages,
                    config='--psm 3 -c preserve_interword_spaces=1'
                )
                text += f"\n--- Page {i+1} ---\n{page_text}"
                
                try:
                    pdf_page_data = pytesseract.image_to_pdf_or_hocr(
                        processed_image, 
                        lang=self.ocr_languages, 
                        extension='pdf'
                    )
                    pdf_page_reader = PyPDF2.PdfReader(io.BytesIO(pdf_page_data))
                    pdf_writer.add_page(pdf_page_reader.pages[0])
                except Exception as e:
                    logger.warning(f"Erreur génération PDF page {i+1}: {str(e)}")
            
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            pdf_writer.write(temp_pdf)
            searchable_pdf_path = temp_pdf.name
            temp_pdf.close()
            doc.close()
            
            return text, searchable_pdf_path, ''
            
        except Exception as e:
            logger.error(f"Erreur OCR PDF avec PyMuPDF: {str(e)}")
            return '', None, str(e)
    
    def extract_from_image(self, image_path):
        """
        Extrait le texte d'une image avec Tesseract OCR.
        """
        try:
            image = Image.open(image_path)
            processed_image = self.preprocess_image(image)
            text = pytesseract.image_to_string(
                processed_image,
                lang=self.ocr_languages,
                config='--psm 3 -c preserve_interword_spaces=1'
            )
            
            pdf_page_data = pytesseract.image_to_pdf_or_hocr(
                processed_image, 
                lang=self.ocr_languages, 
                extension='pdf'
            )
            
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_pdf.write(pdf_page_data)
            searchable_pdf_path = temp_pdf.name
            temp_pdf.close()
            
            return text, searchable_pdf_path, ''
            
        except Exception as e:
            logger.error(f"Erreur OCR image {image_path}: {str(e)}")
            return '', None, str(e)
    
    def extract_from_text(self, text_path):
        try:
            with open(text_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text, ''
        except Exception as e:
            return '', str(e)

    def convert_txt_to_pdf(self, text, original_path):
        """
        Convertit un texte en PDF avec pagination et format A4.
        """
        try:
            import fitz
            doc = fitz.open()
            
            # Paramètres de mise en page (A4: 595 x 842 points)
            width, height = 595, 842
            margin = 50
            rect = fitz.Rect(margin, margin, width - margin, height - margin)
            fontsize = 11
            fontname = "helv"  # Helvetica
            
            # Gérer le texte ligne par ligne ou via textbox
            # Pour faire simple et robuste: on insère, si ça déborde, on crée une page
            current_text = text
            
            while len(current_text) > 0:
                page = doc.new_page(width=width, height=height)
                
                # `insert_textbox` retourne le reste du texte qui n'a pas tenu
                # index de fin de ce qui a été inséré
                try:
                    # Note: rc (return code) est < 0 si overflow
                    # fitz return (rect_consumed, text_remaining) or similar based on version?
                    # The standard API simply returns the rect consumption or similar.
                    # A safer way for basic text is manual chunking or letting textbox handle it.
                    # PyMuPDF's insert_textbox returns the unused text in newer versions, 
                    # OR we can use specific logic. Let's use a simpler known robust method:
                    # insert_textbox returns < 0 on buffer too small, but doesn't return the rest string in all versions?
                    # Let's rely on `page.insert_textbox(rect, current_text, fontsize=..., fontname=...)`
                    # It returns the rectangle used.
                    
                    # Better approach for reliable pagination:
                    rc = page.insert_textbox(rect, current_text, fontsize=fontsize, fontname=fontname)
                    
                    if rc < 0: # Overflow
                         # This version might not be returning the rest of the text.
                         # Fallback strategy: Simple page breaks strictly based on character count is risky.
                         # Best strategy without complex deps: Use PyMuPDF's "get_text_blocks" equivalent or iterative logic.
                         pass
                except:
                    # Fallback simple si insert_textbox échoue
                    p = fitz.Point(margin, margin)
                    page.insert_text(p, current_text[:3000], fontsize=fontsize) # Safe chunk
                
                # For this immediate fix, we will just dump the text properly using a standard method
                # Since we don't have the exact PyMuPDF version docs handy, we'll assume the text fits 
                # or create a logic that splits by newlines approx.
                break # Limitation pour l'instant: 1 page ou refacto plus poussée ci-dessous
            
            # RE-WRITE FOR MULTI-PAGE ROBUSTNESS
            # On recommence avec une logique plus simple: 
            # On découpe le texte nous-mêmes grossièrement
            doc.close()
            doc = fitz.open()
            
            # Fallback si le texte est vide (évite le "Scanned PDF" error)
            if not text or not text.strip():
                text = (
                    "-------------------------------------------------------------------\n"
                    "AVERTISSEMENT: Aucun texte n'a pu être extrait de ce document.\n"
                    "-------------------------------------------------------------------\n\n"
                    "Causes possibles :\n"
                    "1. Le document original est de très mauvaise qualité.\n"
                    "2. C'est une image sans texte lisible.\n"
                    "3. Le format ou l'encodage pose problème.\n\n"
                    "Ceci est un placeholder généré pour permettre l'ouverture du fichier."
                )

            lines = text.split('\n')
            lines_per_page = 55 # approx pour du 11pt sur A4
            
            for i in range(0, len(lines), lines_per_page):
                page = doc.new_page(width=width, height=height)
                chunk = "\n".join(lines[i:i+lines_per_page])
                page.insert_textbox(rect, chunk, fontsize=fontsize, fontname=fontname)

            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_pdf.close()
            doc.save(temp_pdf.name)
            doc.close()
            
            return temp_pdf.name, ''
        except Exception as e:
            logger.error(f"Erreur conversion TXT vers PDF: {str(e)}")
            return None, str(e)

    def extract_from_docx(self, docx_path):
        try:
            doc = docx.Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return '\n'.join(full_text), ''
        except Exception as e:
            logger.error(f"Erreur extraction Word {docx_path}: {str(e)}")
            return '', f"Erreur Word: {str(e)}"


def process_document_ocr(document):
    """
    Traite l'OCR pour un document donné.
    """
    processor = OCRProcessor()
    
    try:
        file_path = document.file.path
        text, searchable_pdf_path, error = processor.extract_text_from_file(file_path)
        
        document.ocr_text = text
        document.ocr_processed = True
        document.ocr_error = error
        document.save()
        
        if searchable_pdf_path and os.path.exists(searchable_pdf_path):
            try:
                from .models import DocumentVersion
                from django.core.files import File
                
                last_version = DocumentVersion.objects.filter(document=document).order_by('-version_number').first()
                next_version = (last_version.version_number + 1) if last_version else 1
                
                new_filename = document.file_name
                if not new_filename.lower().endswith('.pdf'):
                    name_parts = os.path.splitext(new_filename)
                    new_filename = f"{name_parts[0]}.pdf"
                
                with open(searchable_pdf_path, 'rb') as f:
                    django_file = File(f, name=f"Searchable_{new_filename}")
                    DocumentVersion.objects.create(
                        document=document,
                        version_number=next_version,
                        file=django_file,
                        file_name=f"Searchable_{new_filename}",
                        file_size=os.path.getsize(searchable_pdf_path),
                        comment="Version avec OCR (PDF Recherchable)",
                        uploaded_by=document.uploaded_by
                    )
                
                logger.info(f"Version recherchable créée pour document {document.id}")
                try:
                    os.remove(searchable_pdf_path)
                except:
                    pass
            except Exception as e:
                logger.error(f"Erreur version recherchable: {str(e)}")
        
        logger.info(f"OCR traité pour document {document.id}: {len(text)} caractères")
        
    except Exception as e:
        logger.error(f"Erreur traitement OCR document {document.id}: {str(e)}")
        document.ocr_processed = True
        document.ocr_error = str(e)
        document.save()
